"""
Pain_DBS_v9_Pacheco_Rolston.docx = v7 + new pain-motor coupling section.

v7 already contains the v6 -> v8 fixes (time_pos artifact correction +
genetic/biomarker interaction arm). v9 only needs to add:

  1. Methods paragraph: pain-motor coupling analysis (ordinal + logistic
     at cuts, Delta-Delta Spearman with Fisher-z comparison, symmetric
     midpoint anchor for Never-DBS controls).
  2. New Results subsection "Pain-Motor Coupling and Longitudinal
     Decoupling After DBS" (placed after the Genetic/Biomarker subsection,
     before the Network subsection). Matched cohort as primary.
  3. Figure 7 (matched-cohort two-panel figure, main text).
  4. Supplementary Figure S12 (full-cohort sensitivity) + caption.
  5. Supplementary Table S4 caption.
"""
from docx import Document
from docx.shared import Inches
from lxml import etree

SRC = "Pain_DBS_v7_Pacheco_Rolston.docx"
DST = "Pain_DBS_v9_Pacheco_Rolston.docx"

d = Document(SRC)


def replace_paragraph_text(p, new_text):
    ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    runs = p._p.findall(f"{{{ns}}}r")
    for r in runs[1:]:
        p._p.remove(r)
    if runs:
        r0 = runs[0]
        for child in list(r0):
            if child.tag.endswith("}t"):
                r0.remove(child)
        t = etree.SubElement(r0, f"{{{ns}}}t")
        t.text = new_text
        t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    else:
        p.add_run(new_text)


def find_paragraph_index(doc, text_fragment):
    for i, p in enumerate(doc.paragraphs):
        if text_fragment in p.text:
            return i
    return None


# ---------- 1) Methods: pain-motor coupling paragraph ----------
# Insert before the Software paragraph (currently index 48 in v7).
idx_software = find_paragraph_index(d, "All analyses used R version")
assert idx_software is not None, "Software paragraph not found"

methods_coupling = (
    "Pain\u2013motor coupling analyses. To examine whether motor severity "
    "is associated with pain cross-sectionally and whether the pair of "
    "outcomes co-vary within patient, we categorized MDS-UPDRS Part III "
    "into Mild (\u226432), Moderate (33\u201358), and Severe (\u226559) "
    "tiers and dichotomized pain at two pre-specified thresholds (NP1PAIN "
    "\u22651, any pain; NP1PAIN \u22652, moderate or greater). The primary "
    "cross-sectional association was tested by an unadjusted three-tier "
    "ordinal logistic regression (proportional-odds model) in the "
    "propensity-score-matched cohort; because matching balanced age, sex, "
    "disease duration, motor severity, H\u0026Y stage, LEDD, and BMI at "
    "baseline, no further covariate adjustment was applied. Binary logistic "
    "regression at each pain threshold was reported alongside the ordinal "
    "model. Within-patient longitudinal coupling was assessed by Spearman "
    "correlation of \u0394 Pain with \u0394 MDS-UPDRS Part III across the "
    "pre-anchor [\u221224, 0]-month and post-anchor [+6, +18]-month "
    "windows, separately by arm; between-arm differences in correlation "
    "were compared by Fisher-z transformation. Cross-sectional pain\u2013"
    "motor Spearman correlations were also computed in 6-month visit bins "
    "from \u221218 to +24 months relative to anchor. For the coupling "
    "analyses only, Never-DBS controls were re-anchored at the midpoint of "
    "their own follow-up window so that the visit grid was symmetric "
    "pre- and post-anchor across arms; the DBS-anchor (first DBS date) was "
    "retained. Full-cohort patient-anchor sensitivity analyses are "
    "reported in the Supplement."
)

p_soft = d.paragraphs[idx_software]
_ = p_soft.insert_paragraph_before("")
_ = p_soft.insert_paragraph_before(methods_coupling)
print("[1] Methods paragraph inserted")

# ---------- 2) Results: insert subsection before Network section ----------
# Find the Network section body (unique fragment).
idx_net_body = find_paragraph_index(
    d,
    "As an exploratory, post-hoc analysis, we examined whether the multivariate",
)
assert idx_net_body is not None, "Network body paragraph not found"
idx_net_heading = idx_net_body - 1  # the actual section heading

results_heading = "Pain\u2013Motor Coupling and Longitudinal Decoupling After DBS"
results_body = (
    "As a pre-specified exploratory replication of our prior "
    "cross-sectional report (Pacheco-Barrios et al., Parkinson's Dis. "
    "2025;PMID 40003677) showing higher odds of elevated motor severity "
    "among patients with pain, we tested the pain\u2013motor association "
    "at baseline in the propensity-score-matched cohort (primary) and in "
    "the full patient-anchor cohort (sensitivity). In the matched cohort "
    "(n = 149 with complete Pain and MDS-UPDRS Part III data), motor "
    "severity \u226533 at baseline was associated with higher pain tier "
    "(ordinal odds ratio 1.58, 95% CI 0.86\u20132.93, P = .14), in the "
    "same direction as the prior report; the association was similar in "
    "both arms (Never-DBS OR 1.34, DBS OR 1.74; pain \u00D7 arm LRT "
    "P = .63). Full-cohort sensitivity confirmed the association (ordinal "
    "OR 1.76, 95% CI 1.37\u20132.26, P < .001), and a stratum-specific "
    "odds ratio of 3.17 (95% CI 1.21\u20138.30, P = .02) for reporting "
    "any pain among DBS candidates at baseline, consistent with the "
    "channeling pattern noted in Table 1 (Supplementary Figure S12, "
    "Supplementary Table S4). Within-patient longitudinal coupling "
    "diverged by arm: Never-DBS patients showed a positive \u0394 Pain "
    "\u2013 \u0394 MDS-UPDRS Part III Spearman correlation of +0.35 "
    "(P = .005, n = 62), whereas DBS patients showed near-zero coupling "
    "(\u03C1 = +0.11, P = .54, n = 33); the between-arm Fisher-z "
    "comparison was not significant (P = .25) but underpowered given the "
    "DBS sample size. The full-cohort sensitivity reproduced the pattern "
    "(Never-DBS \u03C1 = +0.16, P < .001; DBS \u03C1 = 0.00, P = .99; "
    "Fisher-z P = .30) (Figure 7, panel B). Cross-sectional pain\u2013"
    "motor coupling peaked at the anchor (DBS \u03C1 \u2248 0.36, "
    "Never-DBS \u03C1 \u2248 0.23) and decayed over the following "
    "24 months in both arms. Interpreted cautiously given the exploratory "
    "status and modest sample, these results are consistent with a view "
    "in which stimulation targets motor circuitry and the within-patient "
    "coupling between pain and motor changes is attenuated after DBS, "
    "complementing the earlier observation of strengthened pain\u2013"
    "autonomic network coupling (Figure 6B)."
)

caption_fig7 = (
    "Figure 7. Pain\u2013motor coupling in the propensity-score-matched "
    "cohort. (A) Baseline distribution of MDS-UPDRS Part III severity tiers "
    "(Mild \u226432; Moderate 33\u201358; Severe \u226559) among pain-free "
    "and pain-reporting patients, stratified by arm. (B) Cross-sectional "
    "Spearman correlation between NP1PAIN and MDS-UPDRS Part III at each "
    "6-month visit bin relative to anchor, by arm; marker size encodes n "
    "per bin. Never-DBS controls were re-anchored at the midpoint of their "
    "own follow-up for this analysis so that the visit grid is symmetric "
    "pre- and post-anchor across arms. Dashed vertical line, anchor; "
    "dashed horizontal line, \u03C1 = 0."
)

p_net_heading = d.paragraphs[idx_net_heading]

# Blank spacer → new heading → body → caption → image → blank → (then existing heading)
_ = p_net_heading.insert_paragraph_before("")
_ = p_net_heading.insert_paragraph_before(results_heading)
_ = p_net_heading.insert_paragraph_before(results_body)
_ = p_net_heading.insert_paragraph_before(caption_fig7)
p_img = p_net_heading.insert_paragraph_before("")
run = p_img.add_run()
run.add_picture(
    "outputs/figures/Figure26b_pain_motor_coupling_matched.png",
    width=Inches(6.5),
)
_ = p_net_heading.insert_paragraph_before("")
print("[2] Results subsection + Figure 7 inserted")

# ---------- 3) Supplementary Figure S12 + Supp Table S4 ----------
# Place S12 after S11 (before Supp Table S1 block).
idx_tab_s1 = find_paragraph_index(d, "Supplementary Table S1. Full")
assert idx_tab_s1 is not None

caption_s12 = (
    "Supplementary Figure S12. Pain\u2013motor coupling, full-cohort "
    "patient-anchor sensitivity. Same panels as Figure 7 but computed in "
    "all n = 1,484 PPMI patients with pain observations, using the "
    "patient-specific anchor (first DBS date for cases; first visit for "
    "controls). Baseline ordinal OR for pain tier given motor severity "
    "\u226533: 1.76 (95% CI 1.37\u20132.26, P < .001). Within-patient "
    "\u0394 Pain \u2013 \u0394 MDS-UPDRS Part III Spearman: Never-DBS "
    "\u03C1 = +0.16, P < .001, n = 487; DBS \u03C1 = 0.00, P = .99, n = 50; "
    "Fisher-z between-arm P = .30."
)

caption_t4 = (
    "Supplementary Table S4. Stratum-specific odds ratios for the pain\u2013"
    "motor association at baseline in the matched cohort (primary) and full "
    "patient-anchor cohort (sensitivity). Outcome: NP1PAIN \u22651 or "
    "\u22652 (binary). Predictor: MDS-UPDRS Part III \u226533. Unadjusted "
    "logistic regression; stratum-specific and overall odds ratios with "
    "95% confidence intervals and two-sided P values."
)

p_t1 = d.paragraphs[idx_tab_s1]
_ = p_t1.insert_paragraph_before("")
_ = p_t1.insert_paragraph_before(caption_s12)
p_s12_img = p_t1.insert_paragraph_before("")
p_s12_img.add_run().add_picture(
    "outputs/figures/Figure26_pain_motor_coupling.png",
    width=Inches(6.5),
)
_ = p_t1.insert_paragraph_before("")
_ = p_t1.insert_paragraph_before(caption_t4)
print("[3] Supp Fig S12 + Supp Table S4 inserted")

d.save(DST)
print(f"\nSaved: {DST}")
