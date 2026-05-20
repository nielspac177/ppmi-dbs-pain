"""
Build Pain_DBS_v7_Pacheco_Rolston.docx from v6 with:
  1. Fixed slope numbers in Results paragraph (time_pos artifact correction)
  2. Updated Supp Fig S6 caption (signed calendar time)
  3. Updated Supp Fig S7 caption (signed calendar time, new contrast P=.012)
  4. New Results subsection: "Exploratory Genetic and Biomarker Interactions"
  5. New Methods paragraph describing genetic analyses
  6. New Supplementary Figure S11 (genetics x DBS forest) with embedded image
"""
from docx import Document
from docx.shared import Inches
from copy import deepcopy
from lxml import etree

SRC = 'Pain_DBS_v6_Pacheco_Rolston.docx'
DST = 'Pain_DBS_v8_Pacheco_Rolston.docx'

d = Document(SRC)

# ---------- helpers ----------
def replace_paragraph_text(p, new_text):
    """Replace the visible text of a paragraph while keeping the paragraph XML."""
    # Collect <w:r> runs and drop all but the first.
    ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    runs = p._p.findall(f'{{{ns}}}r')
    for r in runs[1:]:
        p._p.remove(r)
    if runs:
        # Clear existing text elements in the first run, add new w:t
        r0 = runs[0]
        for child in list(r0):
            if child.tag.endswith('}t'):
                r0.remove(child)
        t = etree.SubElement(r0, f'{{{ns}}}t')
        t.text = new_text
        t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    else:
        p.add_run(new_text)

def find_paragraph_index(doc, text_fragment):
    for i, p in enumerate(doc.paragraphs):
        if text_fragment in p.text:
            return i
    return None

# ---------- 1) Results paragraph 62: slope numbers ----------
idx_results = find_paragraph_index(
    d, 'Linear mixed models estimated within-patient Pre-DBS vs. Post-DBS slopes')
assert idx_results is not None, 'Results slope paragraph not found'
new_results_text = (
    "Linear mixed models on signed calendar time from the anchor date estimated "
    "within-patient Pre-DBS vs. Post-DBS slopes of \u22120.011 vs. \u22120.016 "
    "Pain units/month over a \u00B11-year window around surgery (slope contrast "
    "P = .95) and between-arm Post-DBS vs. Never-DBS slopes of +0.004 vs. +0.014 "
    "Pain units/month over 0\u201348 months (slope contrast P = .34) (Table 2, "
    "Supplementary Figure S6). In a wide-window (\u00B15-year) sensitivity fit, "
    "Pre-DBS and Post-DBS slopes were \u22120.009 and +0.007 Pain units/month "
    "(Tukey Pre vs. Post P = .012), consistent with tightly managed preoperative "
    "pain followed by modest postoperative disease progression rather than an "
    "acute surgical reversal (Supplementary Figure S7). A population-averaged "
    "generalized estimating equation (GEE) model with exchangeable working "
    "correlation and inverse-probability weights, adjusted for MDS-UPDRS Part III, "
    "LEDD, body-mass index, sex, GDS, and STAI, yielded concordant null arm \u00D7 "
    "time interactions (Supplementary Table S2). Analytic note: the v6 manuscript "
    "parameterized pre-operative visit time with an unsigned offset, which "
    "inverted the direction of the Pre-DBS slope relative to calendar time; the "
    "present results use signed calendar months and therefore supersede the prior "
    "per-month slope estimates."
)
replace_paragraph_text(d.paragraphs[idx_results], new_results_text)
print(f"[1] Updated Results slope paragraph at index {idx_results}")

# ---------- 2) Supp Fig S6 caption ----------
idx_s6 = find_paragraph_index(d, 'Supplementary Figure S6. Linear mixed-model')
assert idx_s6 is not None
new_s6 = (
    "Supplementary Figure S6. Linear mixed-model predicted Pain trajectories on "
    "signed calendar time from the anchor date. (A) Pre-DBS vs. Post-DBS within "
    "DBS patients over \u00B11 year around surgery; phase-specific slopes "
    "\u22120.011 (Pre-DBS) and \u22120.016 (Post-DBS) Pain units/month, slope "
    "contrast P = .95. (B) Post-DBS vs. Never-DBS over 0\u20134 years; slopes "
    "+0.004 (Post-DBS) and +0.014 (Never-DBS) Pain units/month, slope contrast "
    "P = .34. Model: Pain ~ time \u00D7 phase + (1 + time | patient), IPW-weighted. "
    "This figure replaces the v6 version, which used an unsigned pre-operative "
    "time axis and therefore mis-interpreted the direction of the Pre-DBS slope."
)
replace_paragraph_text(d.paragraphs[idx_s6], new_s6)
print(f"[2] Updated Supp Fig S6 caption at index {idx_s6}")

# ---------- 3) Supp Fig S7 caption ----------
idx_s7 = find_paragraph_index(d, 'Supplementary Figure S7. Wide-window')
assert idx_s7 is not None
new_s7 = (
    "Supplementary Figure S7. Wide-window (\u00B15 years around the anchor date) "
    "linear mixed-model predicted Pain trajectories in the matched cohort on "
    "signed calendar time. Phase-specific slopes: Pre-DBS \u22120.009; Post-DBS "
    "+0.007; Never-DBS +0.001 Pain units/month. Tukey pairwise: Pre vs. Post "
    "P = .012; Pre vs. Never P = .37; Post vs. Never P = .70. The nominal Pre "
    "vs. Post contrast reflects tightly managed preoperative pain (candidates "
    "were typically observed most densely in the weeks before surgery) followed "
    "by modest postoperative disease progression. The v6 version of this figure "
    "used an unsigned pre-operative time axis and reported an inverted pattern."
)
replace_paragraph_text(d.paragraphs[idx_s7], new_s7)
print(f"[3] Updated Supp Fig S7 caption at index {idx_s7}")

# ---------- 4) Supp Table S1 caption refresh ----------
idx_t1 = find_paragraph_index(d, 'Supplementary Table S1. Full linear mixed-effects')
if idx_t1 is not None:
    new_t1 = (
        "Supplementary Table S1. Full linear mixed-effects model coefficients on "
        "signed calendar time from the anchor date: (A) Pre-DBS vs. Post-DBS "
        "within DBS patients (\u00B112 months); (B) Post-DBS vs. Never-DBS "
        "(0\u201348 months, matched cohort n = 170); phase-specific slopes and "
        "between-phase slope contrasts with SE and CI. Available in the online "
        "supplement."
    )
    replace_paragraph_text(d.paragraphs[idx_t1], new_t1)
    print(f"[4] Updated Supp Table S1 caption at index {idx_t1}")

# ---------- 5) Insert Genetics Results subsection before the Network subsection ----------
# Use a phrase unique to the SECTION heading (not the paper title).
idx_network_body = find_paragraph_index(
    d, 'As an exploratory, post-hoc analysis, we examined whether the multivariate')
assert idx_network_body is not None
# Insert before the section *heading*, which is immediately before that body text.
idx_network_heading = idx_network_body - 1

gen_heading = "Exploratory Genetic and Biomarker Interactions"
gen_body = (
    "As a pre-specified exploratory analysis we tested whether baseline genetic "
    "risk or biomarker status modifies the DBS pain response. In the "
    "patient-anchor cohort with Pain observations in both the pre-anchor "
    "[\u221224, 0]-month and post-anchor [+6, +18]-month windows (n = 642; 67 DBS, "
    "575 Never-DBS), we examined four arm \u00D7 stratifier interactions on "
    "\u0394 Pain: (1) a Parkinson-disease polygenic burden score derived as the "
    "allele-count sum across 55 NeuroChip genome-wide association study risk "
    "single-nucleotide polymorphisms, z-standardized and grouped by tertile "
    "(n = 388); (2) APOE-\u03B54 carrier status (0 vs. \u22651 alleles; n = 381); "
    "(3) cerebrospinal-fluid \u03B1-synuclein seeding-amplification assay (SAA) "
    "positivity (n = 574); and (4) GBA-variant carrier status (n = 388, of whom "
    "10 DBS and 34 Never-DBS were GBA carriers). All four arm \u00D7 stratifier "
    "interactions were non-significant (PRS tertile P = .30; APOE-\u03B54 P = .96; "
    "\u03B1-syn SAA P = .71; GBA+ P = .40) (Supplementary Figure S11). Point "
    "estimates were consistent with a uniform DBS effect across strata: the "
    "largest non-significant per-stratum contrast was a numerical DBS benefit "
    "within the middle-PRS tertile (\u0394(DBS \u2212 Never-DBS) = \u22120.41, "
    "95% CI \u22120.95 to +0.13), and a non-significant direction reversal was "
    "observed for GBA carriers (GBA+: +0.16; GBA\u2212: \u22120.14). Clock-gene "
    "polygenic scoring was not feasible because the PPMI NeuroChip array does "
    "not cover CLOCK, BMAL1, PER, CRY, or NR1D variants. These exploratory "
    "interactions, interpreted alongside the primary non-inferiority result, "
    "offer no evidence that common genetic risk or established PD biomarkers "
    "identify a subgroup with materially different longitudinal pain response "
    "to DBS in PPMI; larger samples would be required to rule out modest "
    "modifiers."
)
# Insert before the Network heading
p_heading = d.paragraphs[idx_network_heading]
new_heading_p = p_heading.insert_paragraph_before(gen_heading)
new_body_p = p_heading.insert_paragraph_before(gen_body)
# Blank spacer line
_spacer = p_heading.insert_paragraph_before('')
print(f"[5] Inserted Genetics Results subsection before index {idx_network_heading}")

# ---------- 6) Methods: add a paragraph describing genetics methods ----------
idx_software = find_paragraph_index(d, 'All analyses used R version')
assert idx_software is not None
gen_methods = (
    "Genetic and biomarker \u00D7 DBS interaction analyses. Four pre-specified "
    "exploratory interaction tests on \u0394 Pain were conducted on the "
    "patient-anchor cohort. A Parkinson-disease polygenic burden score was "
    "computed from 55 NeuroChip genome-wide association study risk "
    "single-nucleotide polymorphisms as the allele-dosage sum across variants "
    "(pathogenic variants such as GBA L444P/N370S, LRRK2 G2019S, and SNCA A53T "
    "were excluded); the score was z-standardized across the analytic cohort "
    "and grouped into Low/Middle/High tertiles at the 1/3 and 2/3 empirical "
    "quantiles. APOE-\u03B54 carrier status, GBA-variant carrier status, and "
    "CSF \u03B1-syn SAA results were taken from the PPMI Curated Data Cut "
    "(November 2024) and from the MJF Foundation PPMI database. Each "
    "arm \u00D7 stratifier interaction was assessed as a linear-model F test "
    "on the interaction term; within-stratum DBS vs. Never-DBS contrasts used "
    "Welch two-sample t-tests. Genetic analyses were considered exploratory "
    "and were not adjusted for multiplicity across the four stratifiers."
)
p_soft = d.paragraphs[idx_software]
_spacer2 = p_soft.insert_paragraph_before('')
_new = p_soft.insert_paragraph_before(gen_methods)
print(f"[6] Inserted genetics methods paragraph before index {idx_software}")

# ---------- 7) Supplementary Figure S11 + image before Supp Table S1 ----------
idx_tabS1 = find_paragraph_index(d, 'Supplementary Table S1. Full')
assert idx_tabS1 is not None
s11_caption = (
    "Supplementary Figure S11. Exploratory genetic and biomarker \u00D7 DBS "
    "interaction forest. Point estimate and 95% confidence interval of "
    "\u0394 Pain (DBS \u2212 Never-DBS; post-anchor [+6, +18]-month mean minus "
    "pre-anchor [\u221224, 0]-month mean) within each stratum. Positive values "
    "indicate greater worsening in the DBS arm. Labels at right report n "
    "(DBS / Never-DBS). Interaction P-values (arm \u00D7 stratifier): PRS "
    "tertile P = .30; APOE-\u03B54 P = .96; CSF \u03B1-syn SAA P = .71; "
    "GBA+ P = .40. All four interactions were non-significant, consistent "
    "with a uniform DBS effect across these strata."
)
# Insert caption + image before Supp Table S1
p_tabS1 = d.paragraphs[idx_tabS1]
p_caption = p_tabS1.insert_paragraph_before(s11_caption)
p_image = p_tabS1.insert_paragraph_before('')
run = p_image.add_run()
run.add_picture('outputs/figures/Figure25_genetics_forest.png', width=Inches(6.5))
_spacer3 = p_tabS1.insert_paragraph_before('')
print(f"[7] Inserted Supp Fig S11 + image before index {idx_tabS1}")

d.save(DST)
print(f"\nSaved: {DST}")
