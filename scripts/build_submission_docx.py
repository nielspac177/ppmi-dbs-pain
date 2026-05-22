"""
build_submission_docx.py — JAMA-Neurology-styled .docx generation.

Reads markdown sources from pain_manuscript_220526/0X_*.md and produces
journal-submission-ready .docx files with:

- Times New Roman 12 pt, double-spaced body
- 1-inch margins
- Continuous line numbers in the left margin (JAMA Neurology requirement)
- Page numbers in the footer
- Custom styles for Heading 1/2/3, Abstract section headings, References
- Hanging-indent reference list
- Tables rendered as Word tables (not raw markdown)
- Justified title page block

Usage:
    python3 scripts/build_submission_docx.py

Writes:
    pain_manuscript_220526/01_Title_Page.docx
    pain_manuscript_220526/02_Manuscript_PachecoBarrios_Rolston_2026.docx
    pain_manuscript_220526/03_Supplementary_Materials.docx
    pain_manuscript_220526/04_Cover_Letter.docx
"""
from __future__ import annotations
import re
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement

ROOT = Path(__file__).resolve().parents[1]
BUNDLE = ROOT / "pain_manuscript_220526"
FIG = ROOT / "outputs" / "figures"


# ---------------------------------------------------------------------------
# Document setup helpers
# ---------------------------------------------------------------------------
def setup_document(line_numbers: bool = True, font: str = "Times New Roman",
                   font_size: int = 12, double_spaced: bool = True) -> Document:
    """Build a Document with JAMA-style page setup."""
    doc = Document()

    # Page setup: US Letter, 1-inch margins
    section = doc.sections[0]
    section.page_height = Inches(11)
    section.page_width = Inches(8.5)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    # Continuous line numbers (left margin)
    if line_numbers:
        sect_pr = section._sectPr
        ln_num_type = OxmlElement("w:lnNumType")
        ln_num_type.set(qn("w:countBy"), "1")        # number every line
        ln_num_type.set(qn("w:start"), "1")
        ln_num_type.set(qn("w:restart"), "continuous")
        ln_num_type.set(qn("w:distance"), "200")     # 0.14 inch from text
        sect_pr.append(ln_num_type)

    # Default style: TNR 12pt, double-spaced
    normal = doc.styles["Normal"]
    normal.font.name = font
    normal.font.size = Pt(font_size)
    pf = normal.paragraph_format
    if double_spaced:
        pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    pf.space_after = Pt(0)
    pf.space_before = Pt(0)

    # Headings
    for lvl, size, bold, upper in [(1, 14, True, True),
                                    (2, 13, True, False),
                                    (3, 12, True, False)]:
        try:
            style = doc.styles[f"Heading {lvl}"]
        except KeyError:
            continue
        style.font.name = font
        style.font.size = Pt(size)
        style.font.bold = bold
        style.font.color.rgb = RGBColor(0, 0, 0)
        pf = style.paragraph_format
        pf.space_before = Pt(12)
        pf.space_after = Pt(6)
        pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE if double_spaced else WD_LINE_SPACING.SINGLE
        pf.keep_with_next = True

    # References style (single-spaced, hanging indent)
    if "References" not in [s.name for s in doc.styles]:
        ref = doc.styles.add_style("References", WD_STYLE_TYPE.PARAGRAPH)
        ref.font.name = font
        ref.font.size = Pt(font_size - 1)
        pf = ref.paragraph_format
        pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
        pf.space_after = Pt(6)
        pf.left_indent = Inches(0.5)
        pf.first_line_indent = Inches(-0.5)

    # Footer with page numbers
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.text = ""
    run = fp.add_run()
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = "PAGE"
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run.font.name = font
    run.font.size = Pt(font_size - 1)

    return doc


# ---------------------------------------------------------------------------
# Markdown → DOCX helpers
# ---------------------------------------------------------------------------
_INLINE_BOLD = re.compile(r"\*\*([^*]+)\*\*")
_INLINE_ITAL = re.compile(r"(?<![*_])\*([^*]+)\*(?!\*)")
_INLINE_CODE = re.compile(r"`([^`]+)`")
_INLINE_LINK = re.compile(r"\[([^\]]+)\]\(([^)]+)\)")
_HTML_ENT  = {"&amp;": "&", "&lt;": "<", "&gt;": ">",
              "&nbsp;": " ",
              "&plusmn;": "±", "&minus;": "−", "&mdash;": "—",
              "&ndash;": "–", "&hellip;": "…", "&deg;": "°",
              "&rsquo;": "’", "&lsquo;": "‘",
              "&ldquo;": "“", "&rdquo;": "”"}


def _decode_entities(s: str) -> str:
    for k, v in _HTML_ENT.items():
        s = s.replace(k, v)
    # Hex entities like &#x2014;
    s = re.sub(r"&#x([0-9a-fA-F]+);",
               lambda m: chr(int(m.group(1), 16)), s)
    # Decimal entities
    s = re.sub(r"&#(\d+);",
               lambda m: chr(int(m.group(1))), s)
    # Markdown escapes (must run AFTER inline patterns; here we
    # post-process the literal text just before rendering).
    s = re.sub(r"\\([*_`\[\]()<>#+\-.!\\])", r"\1", s)
    return s


def add_inline(p, text: str) -> None:
    """Append markdown-ish inline runs to paragraph p."""
    text = _decode_entities(text)
    pos = 0
    pattern = re.compile(
        r"\*\*(?P<bold>[^*]+)\*\*"
        r"|(?<![*_])\*(?P<ital>[^*]+)\*(?!\*)"
        r"|`(?P<code>[^`]+)`"
        r"|\[(?P<linktext>[^\]]+)\]\((?P<linkurl>[^)]+)\)"
    )
    for m in pattern.finditer(text):
        if m.start() > pos:
            p.add_run(text[pos:m.start()])
        if m.group("bold"):
            r = p.add_run(m.group("bold"))
            r.bold = True
        elif m.group("ital"):
            r = p.add_run(m.group("ital"))
            r.italic = True
        elif m.group("code"):
            r = p.add_run(m.group("code"))
            r.font.name = "Consolas"
        elif m.group("linktext"):
            # Render link text in body color; preserve URL inline parenthetically.
            r = p.add_run(m.group("linktext"))
            url = m.group("linkurl")
            if url not in m.group("linktext"):
                p.add_run(f" ({url})")
        pos = m.end()
    if pos < len(text):
        p.add_run(text[pos:])


def parse_table_row(line: str) -> list[str]:
    """Parse a markdown table row into stripped cells."""
    parts = [c.strip() for c in line.strip().strip("|").split("|")]
    return parts


def render_markdown(doc: Document, md_text: str,
                    skip_yaml: bool = True,
                    drop_until: str | None = None) -> None:
    """Render markdown into the supplied Document.

    Supports:
      - YAML front matter (stripped)
      - # / ## / ### headings (mapped to Heading 1/2/3)
      - Paragraphs (blank-line separated)
      - **bold**, *italic*, `code`, [link](url) inline
      - Markdown tables (| col | col | with --- header rule)
      - Bullet lists (- item)
      - Numbered lists (1. item)
      - Code fences (```)
      - Horizontal rules (---) — converted to page breaks
    """
    lines = md_text.splitlines()
    i = 0

    # Skip YAML front matter
    if skip_yaml and lines and lines[0].strip() == "---":
        j = 1
        while j < len(lines) and lines[j].strip() != "---":
            j += 1
        i = j + 1

    started = drop_until is None
    while i < len(lines):
        line = lines[i]

        if not started:
            if line.strip().startswith(drop_until or ""):
                started = True
            else:
                i += 1
                continue

        stripped = line.strip()

        # Blank line
        if not stripped:
            i += 1
            continue

        # Horizontal rule → page break
        if stripped == "---":
            doc.add_page_break()
            i += 1
            continue

        # Code fence
        if stripped.startswith("```"):
            j = i + 1
            code_lines = []
            while j < len(lines) and not lines[j].strip().startswith("```"):
                code_lines.append(lines[j])
                j += 1
            p = doc.add_paragraph()
            r = p.add_run("\n".join(code_lines))
            r.font.name = "Consolas"
            r.font.size = Pt(10)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            i = j + 1
            continue

        # Headings
        m = re.match(r"^(#{1,6})\s+(.+?)\s*$", stripped)
        if m:
            level = min(len(m.group(1)), 3)
            p = doc.add_paragraph(style=f"Heading {level}")
            add_inline(p, m.group(2))
            i += 1
            continue

        # Table: header row followed by --- row
        if "|" in line and i + 1 < len(lines) and re.match(
                r"^\s*\|?[\s\-:|]+\|?\s*$", lines[i + 1]):
            header_cells = parse_table_row(line)
            i += 2
            rows = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                rows.append(parse_table_row(lines[i]))
                i += 1
            tbl = doc.add_table(rows=1 + len(rows), cols=len(header_cells))
            tbl.style = "Light Grid Accent 1"
            tbl.autofit = True
            hdr = tbl.rows[0].cells
            for k, hc in enumerate(header_cells):
                hdr[k].text = ""
                p = hdr[k].paragraphs[0]
                p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                r = p.add_run(_decode_entities(hc))
                r.bold = True
                r.font.size = Pt(10)
            for r_idx, row in enumerate(rows):
                cells = tbl.rows[r_idx + 1].cells
                for c_idx, cell_text in enumerate(row):
                    if c_idx >= len(cells):
                        continue
                    cells[c_idx].text = ""
                    p = cells[c_idx].paragraphs[0]
                    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                    add_inline(p, cell_text)
                    for run in p.runs:
                        run.font.size = Pt(10)
            doc.add_paragraph()  # blank line after table
            continue

        # Bullet list
        if stripped.startswith("- ") or stripped.startswith("* "):
            text = stripped[2:]
            p = doc.add_paragraph(style="List Bullet")
            add_inline(p, text)
            i += 1
            continue

        # Numbered list
        m = re.match(r"^(\d+)\.\s+(.+)$", stripped)
        if m:
            text = m.group(2)
            # References-style for the References section
            in_refs = any(p.style.name == "Heading 1"
                          and p.text.strip().lower().startswith("references")
                          for p in doc.paragraphs[-30:])
            if in_refs:
                p = doc.add_paragraph(style="References")
                p.add_run(f"{m.group(1)}. ").bold = False
                add_inline(p, text)
            else:
                p = doc.add_paragraph(style="List Number")
                add_inline(p, text)
            i += 1
            continue

        # Regular paragraph (may span multiple lines until blank).
        # Treat each line as a hard break (one paragraph per line),
        # which suits title pages and address blocks. Within-paragraph
        # markdown soft-wrap is rare in scientific manuscripts.
        para_lines = [line]
        j = i + 1
        while j < len(lines) and lines[j].strip() and \
                not lines[j].strip().startswith(("#", "|", "- ", "* ", "```", "---")) and \
                not re.match(r"^\d+\.\s", lines[j].strip()):
            para_lines.append(lines[j])
            j += 1
        # If all lines share punctuation typical of long-flowing prose
        # (ending in lowercase, mid-clause), join with spaces; otherwise
        # treat each as its own paragraph (typical of address blocks,
        # author lists, key-points blocks).
        looks_like_prose = all(
            not L.strip().endswith((".", "?", "!", ";", ":"))
            and len(L.strip()) > 25
            for L in para_lines[:-1]
        ) if len(para_lines) > 1 else True
        if looks_like_prose:
            paragraph_text = " ".join(l.strip() for l in para_lines)
            p = doc.add_paragraph()
            add_inline(p, paragraph_text)
        else:
            for L in para_lines:
                p = doc.add_paragraph()
                add_inline(p, L.strip())
        i = j


# ---------------------------------------------------------------------------
# Specialized builders
# ---------------------------------------------------------------------------
def build_title_page() -> Path:
    src = (BUNDLE / "01_Title_Page.md").read_text(encoding="utf-8")
    doc = setup_document(line_numbers=False, double_spaced=True)
    render_markdown(doc, src)
    # Centre the very first paragraph after the "Title" heading (the title
    # itself) and the author block.
    centre_next = False
    for para in doc.paragraphs[:20]:
        txt = para.text.strip().lower()
        if txt in {"title", "authors"}:
            centre_next = True
            continue
        if centre_next and txt:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if txt.startswith("deep brain stimulation"):
                # Make title bold + slightly larger
                for r in para.runs:
                    r.bold = True
                    r.font.size = Pt(14)
        if txt in {"corresponding author", "word counts",
                   "manuscript type", "key points", "funding"}:
            centre_next = False
    out = BUNDLE / "01_Title_Page.docx"
    doc.save(out)
    return out


def build_manuscript() -> Path:
    src = (BUNDLE.parent / "manuscript" / "MANUSCRIPT_DRAFT.md").read_text(encoding="utf-8")
    doc = setup_document(line_numbers=True, double_spaced=True)
    # Strip the YAML / header front-matter up to "## Abstract" so the
    # journal-styled doc begins at the abstract.
    abstract_idx = src.find("## Abstract")
    if abstract_idx > 0:
        src = src[abstract_idx:]
    render_markdown(doc, src, skip_yaml=False)
    out = BUNDLE / "02_Manuscript_PachecoBarrios_Rolston_2026.docx"
    doc.save(out)
    return out


def build_supplementary() -> Path:
    src = (BUNDLE / "03_Supplementary_Materials.md").read_text(encoding="utf-8")
    doc = setup_document(line_numbers=False, double_spaced=False)  # supp typically single-spaced
    # Set Normal to single-spaced for supplement
    for s_name in ["Normal", "Heading 1", "Heading 2", "Heading 3"]:
        try:
            doc.styles[s_name].paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        except KeyError:
            pass
    render_markdown(doc, src)
    out = BUNDLE / "03_Supplementary_Materials.docx"
    doc.save(out)
    return out


def build_cover_letter() -> Path:
    src = (BUNDLE / "04_Cover_Letter.md").read_text(encoding="utf-8")
    doc = setup_document(line_numbers=False, double_spaced=False)
    for s_name in ["Normal", "Heading 1", "Heading 2", "Heading 3"]:
        try:
            doc.styles[s_name].paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        except KeyError:
            pass
    # Cover letter: 1.15 line spacing, left-aligned
    for sty in doc.styles:
        if sty.type == WD_STYLE_TYPE.PARAGRAPH and sty.paragraph_format is not None:
            sty.paragraph_format.space_after = Pt(8)
    render_markdown(doc, src)
    out = BUNDLE / "04_Cover_Letter.docx"
    doc.save(out)
    return out


# ---------------------------------------------------------------------------
def main() -> None:
    for builder, label in [
        (build_title_page, "Title page"),
        (build_manuscript, "Manuscript"),
        (build_supplementary, "Supplementary"),
        (build_cover_letter, "Cover letter"),
    ]:
        out = builder()
        size = out.stat().st_size
        print(f"[OK] {label:13s} → {out.name}  ({size // 1024} KB)")


if __name__ == "__main__":
    main()
