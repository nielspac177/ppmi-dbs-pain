"""
Fold the methods schematic into Pain_DBS_v9_Pacheco_Rolston.docx:

 1. Renumber every "Figure N" / "Figure NA" / "Figure NB" in the body
    according to the final desired numbering:

       Methods    -> Figure 1   (new)
       STROBE     -> Figure 2   (was Figure 1)
       Motor PC   -> Figure 3   (was Figure 2)
       Landmark   -> Figure 4   (was Figure 3)
       Alluvial   -> Figure 5   (was Figure 4)
       Stratified -> Figure 6   (was Figure 5)
       Pain-motor -> Figure 7   (was Figure 7 -- stays; doc order fix)
       Network    -> Figure 8A/B (was Figure 6A/6B -- doc order fix)

 2. Insert the Methods schematic image + caption right after the
    "Methods" section heading.

 3. Save back to Pain_DBS_v9_Pacheco_Rolston.docx (in-place update).
"""
import re
from docx import Document
from docx.shared import Inches
from lxml import etree

DOC = "Pain_DBS_v9_Pacheco_Rolston.docx"
FIG = "outputs/figures/Figure_methods_schematic.png"

# old caption number (integer or string '6A'/'6B') -> new caption number string
NUMBER_MAP = {
    "1":  "2",
    "2":  "3",
    "3":  "4",
    "4":  "5",
    "5":  "6",
    "6":  "8",     # Network standalone ref
    "6A": "8A",
    "6B": "8B",
    "7":  "7",     # Pain-motor stays
}

FIGURE_RE = re.compile(r"\bFigure (\d+)(A|B)?\b")


def bump_figures(text: str) -> str:
    def repl(m):
        digits, letter = m.group(1), m.group(2) or ""
        key = digits + letter
        if key in NUMBER_MAP:
            return f"Figure {NUMBER_MAP[key]}"
        # Fallback: bump numeric part by +1 if not in map
        return f"Figure {int(digits) + 1}{letter}"
    return FIGURE_RE.sub(repl, text)


def replace_paragraph_text(p, new_text: str) -> None:
    ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    runs = p._p.findall(f"{{{ns}}}r")
    for r in runs[1:]:
        p._p.remove(r)
    if runs:
        r0 = runs[0]
        for child in list(r0):
            if child.tag.endswith("}t"):
                r0.remove(child)
        t = etree.SubElement(r0, f"{{{ns}}}t")
        t.text = new_text
        t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    else:
        p.add_run(new_text)


d = Document(DOC)

# ---------- 1. Renumber every Figure N reference in paragraphs ----------
renum_count = 0
for p in d.paragraphs:
    old = p.text
    if "Figure " not in old:
        continue
    new = bump_figures(old)
    if new != old:
        replace_paragraph_text(p, new)
        renum_count += 1
print(f"[1] Renumbered Figure references in {renum_count} paragraphs")

# ---------- 2. Insert Methods Figure 1 right after "Methods" heading ----
# Re-find the Methods heading in the freshly-renumbered doc.
idx_methods = None
for i, p in enumerate(d.paragraphs):
    if p.text.strip() == "Methods":
        idx_methods = i
        break
assert idx_methods is not None, "Methods heading not found"

caption = (
    "Figure 1. Study design, anchor alignment, time windows, and analysis "
    "tiers. From the Parkinson's Progression Markers Initiative (PPMI) "
    "Curated Data Cut (November 2024), an analytic cohort of 1,484 "
    "idiopathic Parkinson-disease patients with at least one NP1PAIN "
    "observation was assembled (105 DBS, 1,379 Never-DBS). DBS patients "
    "were anchored at the first DBS surgery date (a clinical event); "
    "Never-DBS patients were anchored at the midpoint of their own "
    "follow-up so that each patient contributed symmetric pre- and "
    "post-anchor bins (this is a geometric device rather than a clinical "
    "event). Time windows: pre-anchor baseline [\u221224, 0] months, "
    "primary post window [+6, +18] months, and landmark assessments at "
    "6/12/18/24/36/48 months \u00B1 6. Analyses are organised in four "
    "tiers. Primary: landmark \u0394 Pain and two-one-sided-tests (TOST) "
    "non-inferiority at a \u00B11-point margin. Secondary: linear mixed "
    "models (\u00B11 year and 0\u201348 months), generalised estimating "
    "equations (IPW-weighted), Kaplan\u2013Meier time to pain worsening, "
    "and a graphical-LASSO partial-correlation network over 15 non-motor "
    "nodes. Sensitivity: 1:2 nearest-neighbour propensity-score matching "
    "on age, sex, disease duration, MDS-UPDRS Part III, Hoehn \u0026 Yahr "
    "stage, LEDD, and body-mass index (caliper 0.2), n = 170; baseline "
    "covariates were taken from the visit closest to the anchor within "
    "[\u221224, 0] months. Exploratory: genetic and biomarker \u00D7 DBS "
    "interactions and pain\u2013motor coupling."
)

# To insert *after* the heading, insert immediately before the paragraph
# that follows the heading. Use insert_paragraph_before on paragraph[idx+1].
anchor = d.paragraphs[idx_methods + 1]
p_spacer1 = anchor.insert_paragraph_before("")
p_img = anchor.insert_paragraph_before("")
p_img.add_run().add_picture(FIG, width=Inches(6.5))
p_cap = anchor.insert_paragraph_before(caption)
p_spacer2 = anchor.insert_paragraph_before("")
print(f"[2] Methods Figure 1 inserted after Methods heading (idx {idx_methods})")

d.save(DOC)
print(f"\nSaved in place: {DOC}")
